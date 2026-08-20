from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


OUT = Path(r"D:\Family\materials\榜样教育\榜样教育项目纲领与五份PPT逐页解读.docx")

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172033"
MUTED = "55657A"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF7E6"
GOLD = "9A6700"
RED = "9B1C1C"
GREEN = "176B4D"
WHITE = "FFFFFF"
BORDER = "D7DEE8"


def set_font(run, latin="Calibri", east_asia="Microsoft YaHei", size=None,
             color=None, bold=None, italic=None):
    run.font.name = latin
    if run._element.get_or_add_rPr().rFonts is None:
        run._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
    rfonts = run._element.get_or_add_rPr().rFonts
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_paragraph_border(paragraph, color=BORDER, size="8", space="3"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])


def add_para(doc, text="", style=None, size=None, color=INK, bold=False,
             italic=False, align=None, before=0, after=6, line=1.10,
             keep_with_next=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep_with_next
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, label_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(label + "：")
    set_font(r, size=11, color=label_color, bold=True)
    r = p.add_run(text)
    set_font(r, size=11, color=INK)
    add_para(doc, "", after=6)


def add_labeled_block(doc, labels):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.10
    for idx, (label, text, color) in enumerate(labels):
        if idx:
            p.add_run().add_break()
        r = p.add_run(label + "：")
        set_font(r, size=11, color=color, bold=True)
        r = p.add_run(text)
        set_font(r, size=11, color=INK)
    return p


def add_slide_entry(doc, number, title, intent, reading, judgment, directive):
    h = doc.add_paragraph(style="Heading 3")
    h.paragraph_format.keep_with_next = True
    r = h.add_run(f"第{number:02d}页｜{title}")
    set_font(r, size=12, color=DARK_BLUE, bold=True)
    add_labeled_block(doc, [
        ("页面主旨", intent, BLUE),
        ("战略解读", reading, GREEN),
        ("关键判断", judgment, GOLD),
        ("纲领转化", directive, RED),
    ])


def add_major_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.page_break_before = True
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=16, color=BLUE, bold=True)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=13, color=BLUE, bold=True)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=12, color=DARK_BLUE, bold=True)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
section.different_first_page_header_footer = True

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    s = styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run("榜样教育项目纲领与材料解读")
set_font(hr, size=8.5, color=MUTED, bold=True)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = fp.add_run("内部讨论稿  |  ")
set_font(fr, size=8.5, color=MUTED)
add_field(fp, "PAGE")

# Cover: editorial_cover pattern, with standard_business_brief body tokens.
for _ in range(5):
    add_para(doc, "", after=12)
add_para(doc, "STRATEGY CHARTER · 2026", size=10, color=GOLD, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
add_para(doc, "榜样教育项目纲领", size=30, color=NAVY, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
add_para(doc, "五份PPT逐页解读与战略汇总", size=15, color=DARK_BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
add_para(doc, "从课程公司到家庭成长数据科技与服务平台",
         size=11, color=MUTED, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=72)
add_para(doc, "内部讨论稿 · 仅用于战略对齐，不构成法律、税务或投资意见",
         size=9.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
add_para(doc, "2026年8月9日", size=11, color=NAVY, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
doc.add_page_break()

add_h2(doc, "编制说明")
add_para(doc, "本文件基于用户提供的五份PPT进行逐页阅读、跨材料比对和战略提炼。文中把原材料的陈述视为内部方案或待验证假设，不把其中的市场规模、估值、上市、转化率和增长目标自动视为已被外部证据证明的事实。")
add_callout(doc, "使用方式", "前半部分用于管理层统一方向、原则、边界和未来12个月行动；后半部分保留逐页解读，便于追溯每一项纲领结论来自哪一页材料。", PALE_BLUE, BLUE)

add_h2(doc, "五份材料登记")
source_rows = [
    ("A", "榜样教育战略白皮书_30页演讲汇报版_D盘版", "30", "战略主干；与微信版30页画面逐页相同"),
    ("B", "榜样教育战略白皮书_30页演讲汇报版_微信版", "30", "战略主干；二进制不同但导出画面一致"),
    ("C", "家庭教育大模型平台科技公司项目合作方案", "10", "合作愿景、技术路径、组织与资本叙事"),
    ("D", "榜样教育新商业模式对外宣发PPT_原图版(2)", "20", "商业模式、产品界面、裂变与用户经营"),
    ("E", "榜样科技创业合伙人股权架构设计(1)", "14", "股权、人才激励、控制权与退出机制"),
]
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
set_table_geometry(table, [600, 4200, 800, 3760])
headers = ["编号", "材料", "页数", "在本文件中的角色"]
for i, htxt in enumerate(headers):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, NAVY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(htxt)
    set_font(r, size=10, color=WHITE, bold=True)
set_repeat_table_header(table.rows[0])
for row in source_rows:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        p = cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(value)
        set_font(r, size=9.5, color=INK, bold=(i == 0))
set_table_geometry(table, [600, 4200, 800, 3760])
add_para(doc, "来源完整性说明：材料D的PPT包可提取20页XML、文本和7张嵌入原图，但本机PowerPoint报告“文件或目录损坏且无法读取”。本文件已按包内20页顺序完成解读；源文件保持原样归档，建议另行修复后再作为正式对外演示文件。", size=9.5, color=RED, after=8)

add_h2(doc, "阅读导航")
add_labeled_block(doc, [
    ("第一部分", "项目纲领：使命、定位、产品、增长、数据AI、组织、治理与指标。", BLUE),
    ("第二部分", "战略冲突与裁决：对五份材料中相互矛盾或风险过高的表述作统一判断。", GREEN),
    ("第三部分", "未来12个月：从概念叙事转入产品验证、交付标准化和治理建设。", GOLD),
    ("第四部分", "逐页解读：覆盖五个文件；两份视觉完全相同的白皮书采用共同30页解读并双版本映射。", RED),
])

add_major_heading(doc, "第一部分  项目纲领")
add_callout(doc, "一页结论", "榜样教育不应把自己定义为AI工具公司，也不应急于宣布自己已经是平台。更稳健的路径是：以家庭真实改变为价值，以标准化交付为底座，以最小必要且获授权的成长数据为资产，以AI协同提升效率和个性化，以会员与长期服务形成复购，最终在结果可验证、单位经济模型成立后再开放生态。", PALE_GOLD, GOLD)

charter_sections = [
    ("1. 使命", "帮助父母与孩子在真实家庭场景中形成可见、可衡量、可持续的积极改变，并让高质量家庭成长服务能够以更低边际成本被稳定复制。使命的对象不是“数据”或“AI”，而是家庭关系、行为与成长结果。"),
    ("2. 战略定位", "榜样教育定位为“家庭成长数据科技与服务平台的建设者”。在当前阶段，经营实体首先仍是一家产品与服务公司；只有当复购产品、标准化交付、持续授权的数据和可扩展供给同时成立时，平台属性才真正成立。"),
    ("3. 核心用户", "决策者与付费者主要是父母，成长对象包括孩子与家庭关系。产品入口可以是亲子沟通、学习习惯、手机管理、自驱力和情绪等高频问题，但表达必须避免污名化儿童、制造焦虑或承诺确定疗效。"),
    ("4. 价值主张", "用户购买的不是大模型，而是更清楚的问题判断、更容易执行的行动计划、持续的反馈支持、可见的阶段变化以及可信任的长期陪伴。AI属于价值实现机制，不是价值本身。"),
    ("5. 战略顺序", "先选择一个高频问题做出高留存产品闭环，再把课程、助教、评估、运营和复购标准化；随后设计数据结构与授权治理，将AI嵌入任务、反馈和预警；最后才扩展会员、城市服务、专家市场与生态合作。"),
    ("6. 核心产品阶梯", "建议采用“家庭测评/内容入口—21天低门槛挑战—90天家庭成长计划—年度会员—咨询、沙龙与城市服务”的阶梯。第一阶段只验证“家长训练营+AI陪练+成长报告”这一核心组合，不同时建设所有界面和生态模块。"),
    ("7. 交付原则", "交付必须同时覆盖课程内容、每日任务、陪伴反馈、阶段评估、成果呈现和复购建议。关键流程形成SOP，但SOP不能替代专业判断；涉及儿童风险、心理危机、虐待或自伤线索时，必须有人工升级和专业转介机制。"),
    ("8. 增长原则", "增长围绕真实结果而非流量投放。内容与直播建立认知和信任，体验产品筛选意愿，社群交付产生结果，成果展示推动续费和自愿推荐。裂变必须单层、透明、可追溯，奖励以成长权益为主，避免多层返佣和“拉人头”观感。"),
    ("9. 商业模式", "收入结构由课程现金流逐步升级为会员订阅、AI陪伴、咨询活动和生态服务的复合收入。平台估值不是目标函数；优先验证毛利、留存、续费、获客回收期和家庭生命周期价值。材料中的1000万、1亿、10亿、100亿属于阶段愿景，须由可审计经营模型支撑。"),
    ("10. 数据与AI", "家庭成长数据的价值来自在明确目的、最小必要、知情授权、分级访问和可撤回机制下改善服务，而不是把家庭隐私本身视为可出售资产。AI应承担问答、任务提醒、总结、初步建议和服务辅助；高风险判断、关键反馈和未成年人相关决策保留人工责任。"),
    ("11. 组织架构", "设置产品中心、交付中心、数据AI中心三个价值主链条；增长与商业化作为横向能力协同。优先补齐产品负责人、交付负责人和数据产品经理，再补AI应用和增长负责人。组织规模随闭环证据扩张，不以人数替代能力。"),
    ("12. 公司与治理", "独立新公司可以隔离风险、明确资产和吸引人才，但前提是先确定知识产权、客户数据、品牌、课程资产、人员与现金流的归属。重大事项、关联交易、数据安全、未成年人保护、内容审核和服务质量应进入董事会或治理清单。"),
    ("13. 人才与股权", "股权是长期合约而非即时福利。可采用职业经理人、事业伙伴、创业合伙人分层机制，并配合4年成熟、悬崖期、业绩条件和退出回购。具体比例、回购价格、控制权和税务处理必须在授予前经法律、财务与税务专项审查。"),
    ("14. 北极星指标", "建议北极星指标为“持续获得可验证成长结果的有效家庭数”。支持指标包括：测评完成率、21天完成率、90天结果率、阶段续费率、会员续费率、主动推荐率、自然新增占比、服务升级率、人工升级响应时效以及必要数据字段的授权覆盖率。"),
]
for title, text in charter_sections:
    add_h2(doc, title)
    add_para(doc, text)

doc.add_page_break()
add_h2(doc, "15. 经营红线")
redlines = [
    ("未经产品闭环验证，不扩大功能、城市或赛道。", "平台化必须晚于结果验证。"),
    ("不把“AI”“大模型”“数据资产”作为对家长的第一价值承诺。", "品牌第一句应是家庭改变。"),
    ("不把儿童与家庭隐私当作可交易资产。", "数据价值来自受托使用和服务改进。"),
    ("不以多层返佣、强推焦虑或夸大疗效驱动增长。", "裂变以自愿、透明和真实结果为前提。"),
    ("不把估值、上市或百亿收入当作已经成立的经营事实。", "所有外部叙事必须与可验证证据匹配。"),
    ("不在法律、税务、估值和退出规则未明确前授予实质权益。", "股权先制度、后授予。"),
]
for i, (rule, reason) in enumerate(redlines, 1):
    add_h3(doc, f"红线{i}  {rule}")
    add_para(doc, reason, color=MUTED)

add_major_heading(doc, "第二部分  战略冲突与统一裁决")
conflicts = [
    ("“黑灯工厂式全自动化”与“家庭教育依赖信任和专业服务”冲突", "裁决：将目标改为“人机协同的高自动化运营系统”。自动化用于触达、排程、提醒、总结和常规问答；诊断、危机、价值判断、投诉和关键成长反馈保留人工责任。"),
    ("“用户数据资产可调用、支撑高估值”与家庭隐私和未成年人保护冲突", "裁决：把数据定义为受托管理的服务资产，不是可自由交易资产。数据采集必须目的限定、最小必要、分级授权、可撤回、可删除，并建立未成年人和敏感信息的更高保护级别。"),
    ("“替代传统教培”与现阶段依赖课程现金流冲突", "裁决：表述为“从课程交易升级为长期家庭成长服务”。课程不是要被立即取消，而是作为入口和现金流，被新的陪伴、会员和数据闭环重新组织。"),
    ("“跨赛道快速复用”与家庭教育的专业深度冲突", "裁决：前三年不把跨赛道复制作为主任务。底层技术可以复用，但知识、服务流程、风控和供给不能仅靠替换知识库复制。"),
    ("“一年完成核心闭环、100万家庭、50人、3000万投入”之间资源约束不清", "裁决：采用阶段闸门。先以一个场景、一个核心产品、一个城市或私域样本验证结果与单位经济；达到闸门后再扩团队和投入。"),
    ("“经济收益不等于决策权”与合伙人长期信任之间可能失衡", "裁决：创始人控制权可以被保护，但同时需要信息权、经济权、重大利益冲突回避、稀释规则和退出价格等少数权益保护，避免激励计划沦为单方裁量。"),
    ("“放弃现金×创业系数”简洁，但缺少估值和公平基础", "裁决：该公式仅可作为谈判提示，不可直接作为授予定价。最终方案应基于岗位价值、公司估值、授予工具、成熟条件、税负、稀释和离职情形进行场景测算。"),
]
for i, (title, decision) in enumerate(conflicts, 1):
    add_h2(doc, f"冲突{i}  {title}")
    add_para(doc, decision)

add_major_heading(doc, "第三部分  未来12个月行动纲要")
add_para(doc, "未来12个月的目标不是证明公司已经是平台，而是证明一个家庭成长产品能够稳定产生结果、持续复购、沉淀合规数据并被团队复制。")

milestones = [
    ("0-30天", "确定单一问题与核心客群；冻结MVP范围；定义结果指标、数据字典和授权文本；明确新公司与原业务资产边界。", "产品假设、交付蓝图、数据治理清单、合作分工RACI"),
    ("31-90天", "用现有SaaS底座上线最小闭环；完成家长训练营、AI陪练和成长报告；招募首批种子家庭。", "可用产品、SOP v1、首批真实交付与问题清单"),
    ("4-6个月", "连续复盘完成率、结果率、续费和人工成本；修订画像与Agent；形成高风险人工升级机制。", "结果报告、单位经济模型、SOP v2、AI安全评估"),
    ("7-9个月", "验证内容+私域+转介绍增长；试点一个城市沙龙或机构合作；不做全国扩张。", "增长实验台账、单城/单机构试点结论"),
    ("10-12个月", "根据闸门决定扩大、收缩或转向；完成正式治理与人才激励文件；形成下一年度预算。", "董事会决策包、年度经营计划、经审查的股权激励方案"),
]
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
set_table_geometry(table, [1200, 5000, 3160])
for i, htxt in enumerate(("阶段", "关键任务", "阶段交付")):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, NAVY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(htxt)
    set_font(r, size=10, color=WHITE, bold=True)
set_repeat_table_header(table.rows[0])
for row in milestones:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        p = cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(value)
        set_font(r, size=9.5, color=INK, bold=(i == 0))
set_table_geometry(table, [1200, 5000, 3160])

doc.add_page_break()
add_h2(doc, "阶段闸门")
gates = [
    ("产品闸门", "核心场景的任务完成率与阶段结果率连续改善，且不存在不可接受的安全或投诉问题。"),
    ("商业闸门", "获客、交付和服务成本可测；续费与生命周期价值能够覆盖合理获客和人工成本。"),
    ("数据闸门", "必要字段定义清楚、授权可证明、访问可审计、删除可执行，且模型效果确实因数据改善。"),
    ("组织闸门", "核心SOP在非创始团队成员手中也能稳定执行，关键岗位责任和升级机制清晰。"),
    ("扩张闸门", "单一产品和单一渠道先成立，再开放第二场景、第二城市或第三方生态。"),
]
for title, text in gates:
    add_h3(doc, title)
    add_para(doc, text)

add_major_heading(doc, "第四部分  逐页解读")
add_callout(doc, "版本映射", "材料A（D盘版）与材料B（微信版）虽文件哈希不同，但30页导出画面逐页完全相同。因此下列30页逐页解读同时适用于两个文件，相当于覆盖两份材料的60个物理页。", PALE_BLUE, BLUE)

add_h2(doc, "4.1 战略白皮书两版本（共同30页解读）")
whitepaper = [
    (1, "封面：AI成长平台", "确立“家庭成长数据科技公司”而非AI工具公司的主叙事。", "这是整套材料最重要的定位句，决定后续产品、组织、融资和数据故事。", "定位清楚，但“平台”应作为建设目标而非当前既成事实。", "正式口径：我们正在建设以家庭改变为核心的成长数据科技与服务平台。"),
    (2, "汇报主线", "用机会、定位、模式、能力、路径五个问题组织30页内容。", "这是一张管理层对齐地图，适合成为董事会讨论的目录。", "导出画面中标题与副标题重叠，影响正式使用。", "保留五问框架，并修复版式；每一问配置一项可验证决策。"),
    (3, "机会判断", "核心机会是重构家庭成长服务，课程与AI分别是入口和能力。", "把价值中心从技术转向长期服务，是正确的战略起点。", "仍缺少目标家庭、核心问题和效果证据。", "将机会定义为“可持续家庭改变的供给缺口”，并用用户研究验证。"),
    (4, "需求换挡", "指出学习焦虑、关系焦虑和陪伴缺口三类需求。", "把市场从提分扩展到家庭系统，有助于构建更长生命周期。", "三类需求跨度较大，第一阶段不能同时承接。", "选择一个高频问题作为MVP，其余问题进入产品路线图而非首发范围。"),
    (5, "用户购买确定性", "强调用户购买的是孩子改变和关系改善，而不是AI。", "该页应成为品牌、产品和销售的共同价值原则。", "家庭成长不是绝对可控结果，不能承诺“确定疗效”。", "承诺可执行过程、可见阶段进步和可信陪伴，不承诺必然结果。"),
    (6, "先闭环后平台", "给出课程—改变—记录—AI—数据库—生态的战略顺序。", "因果链完整，能够抑制“先搭大平台”的冲动。", "每一节点缺少过关指标，容易又变成概念链。", "为每一环设置闸门：完成率、结果率、续费、数据授权、模型增益和供给质量。"),
    (7, "公司定位升级", "用“不是课程公司、不是聊天工具、而是平台公司”形成三段式定位。", "适合内部澄清边界，也适合融资叙事。", "过早否定课程公司可能削弱现有现金流与能力基础。", "表述为“以课程能力为起点，升级为家庭成长数据科技与服务公司”。"),
    (8, "复合收入", "提出入口、核心、订阅、数据、生态五层产品与价格带。", "产品阶梯能够连接获客、结果、复购和高客单服务。", "价格只是假设；“数据产品”不能等同出售画像。", "用试点验证价格和转化；数据报告作为会员服务与机构洞察，需严格授权。"),
    (9, "成长旅程", "把内容、直播、体验课、训练营、AI陪伴和会员串成用户旅程。", "从漏斗升级为旅程，有利于逐节点设计任务和数据。", "页面底部误用了通用平台说明，与本页重点不完全匹配。", "为每一节点明确用户任务、退出条件、人工责任、数据采集和下一步推荐。"),
    (10, "增长飞轮", "增长围绕学习、改变、数据、AI和更好体验循环。", "飞轮把裂变建立在结果而非投放上，方向正确。", "“数据越多越准”不是必然，还取决于质量、偏差和治理。", "将飞轮改为“高质量授权数据—更好建议—更好结果—更多信任家庭”。"),
    (11, "产品矩阵", "覆盖父母、孩子、家庭关系、AI陪伴和数据产品。", "多角色视角符合家庭系统的真实复杂性。", "同时服务全家会显著增加内容、安全和交付难度。", "第一年以父母端为主，孩子端只做低风险任务与反馈，并设置家长知情和人工审核。"),
    (12, "爆款组合", "建议家长训练营、AI陪练、成长报告作为第一阶段组合。", "这是五份材料中最可执行的MVP定义。", "仍需限定单一问题、目标年龄、周期和成功标准。", "用一个问题场景建立可复用闭环，再扩展到第二问题。"),
    (13, "复制型交付", "将课程、助教、运营、评估、复购和质检流程化。", "该页把组织能力放在技术之前，符合平台成长规律。", "底部通用说明重复，且SOP可能被误解为机械服务。", "SOP标准化常规动作，专业判断与危机处理建立人工升级。"),
    (14, "交付指标", "用完课率、打卡率、画像完整度、活跃率和续费率管理交付。", "开始把价值主张转成责任和指标，是从PPT走向经营的关键。", "缺少真实结果、安全、投诉和人工成本指标。", "增加90天结果率、投诉率、风险升级时效、单家庭交付工时和净推荐意愿。"),
    (15, "AI的正确位置", "AI放在交付闭环，而不是品牌第一句。", "这是项目最重要的技术原则。", "需要进一步明确AI可做与不可做的边界。", "AI负责建议与辅助，专业诊断、危机处置和关键未成年人决策由人负责。"),
    (16, "AI平台架构", "知识库、用户画像、成长规划、陪练、督导和数据中台形成链路。", "架构从用户任务出发，比堆模型名称更有价值。", "底部通用平台说明再次重复；缺少权限、审计和模型评估层。", "在架构中加入同意管理、访问控制、人工升级、安全评测和数据生命周期。"),
    (17, "Agent体系", "按家长、孩子、助教、家庭和管理层配置Agent。", "任务与输出物绑定，是可评估Agent的正确设计方式。", "“孩子陪练”需要年龄适配、内容安全和家长可见性。", "每个Agent必须有用户任务、允许动作、禁区、人工责任人和质量指标。"),
    (18, "家庭成长数据库", "把父母、孩子、互动和改变路径视为护城河。", "长期纵向数据确实可能提升个性化和研究能力。", "护城河不能建立在过度采集或用户难以退出之上。", "把护城河定义为“信任+高质量授权数据+可验证交付方法”，而非数据占有量。"),
    (19, "数据采集结构", "提出父母、孩子、家庭三类画像。", "从第一天设计结构能避免后续只有聊天记录而无可用数据。", "画像字段容易越界，尤其是情绪、家庭冲突和儿童信息。", "先做必要性评估、敏感等级、保存期限和删除机制，再采集；默认最小化。"),
    (20, "增长联动", "内容、直播、社群、成果、推荐和会员形成增长链。", "增长与交付连接，避免把获客和服务割裂。", "导出画面标题和副标题重叠，底部说明也与本页重复；“成果展示”需隐私边界。", "修复版式；案例默认匿名与单独授权，儿童影像和家庭细节不得被默认营销使用。"),
    (21, "用户质量指标", "提出有效线索、体验转化、任务完成、续费和转介绍目标。", "从线索数量转向用户质量，是保护交付的正确变化。", "30%、15-25%、70%、30%、20%均为未验证基准。", "将数字标记为试点假设，连续三个周期复盘后再成为正式预算指标。"),
    (22, "组织重建", "产品中心定义价值，交付中心兑现价值，数据AI中心放大价值。", "三中心结构清晰，适合早期公司避免技术主导一切。", "缺少增长、合规、财务和跨中心决策机制。", "以三中心为主链，增长与风控横向协同；建立周经营会和月度治理复盘。"),
    (23, "招聘优先级", "先招产品、交付、数据，再补AI和增长。", "排序与战略顺序一致，能避免先堆研发。", "岗位产出清楚但选人标准、预算和创始人兼任期未说明。", "每个岗位绑定90天交付；未过产品闸门前由小团队兼任而非扩编。"),
    (24, "财务阶段", "以1000万、1亿、10亿、100亿描绘收入结构升级。", "强调收入质量而非规模，方向正确。", "数字跨度巨大，未给市场、价格、留存、毛利和资本需求模型。", "把四个数字降级为远景情景，正式财务模型从家庭数×ARPU×留存×毛利构建。"),
    (25, "董事会指标", "从用户、产品、商业、数据、组织五维判断战略。", "这是纲领落地为治理的关键页。", "“画像完整度”容易诱导过度采集，且缺少安全和现金指标。", "改为“必要字段授权覆盖率”，并增加现金跑道、重大投诉、风险事件和单位交付成本。"),
    (26, "融资故事", "以市场、复合模式和数据壁垒替代单纯技术概念。", "投资叙事从模型能力转向经营确定性是正确方向。", "仍缺少留存、毛利、CAC回收期和证据样本。", "融资前先形成一套可审计的产品结果与单位经济证据包。"),
    (27, "三年路线", "产品验证、规模复制、平台生态三年递进。", "每年打一场主仗，节奏优于并行扩张。", "10万和100万家庭目标需与交付能力、获客成本、数据治理匹配。", "每年目标以闸门定义；未达闸门则延后平台开放，不以年份倒逼冒进。"),
    (28, "第一年五件事", "聚焦爆款产品、交付SOP、AI助手、数据体系和增长飞轮。", "五件事覆盖价值、交付、技术、资产和增长，结构完整。", "负责人名称存在“中心先于人才”的风险，验收标准不够量化。", "为五件事设单一负责人、季度里程碑、预算和停止条件。"),
    (29, "最终共识", "产品标准化、数据资产化、AI协同化、生态平台化。", "四化顺序能够成为公司战略口号。", "“数据资产化”需避免被误解为数据交易。", "正式表述改为“产品标准化、数据可信化、AI协同化、生态平台化”。"),
    (30, "下一步", "要求从统一共识转入产品体系、PRD、融资BP和组织手册。", "明确“停止扩写、启动验证”，是材料中最具执行力的结尾。", "四类文件若同时启动会分散核心团队。", "优先完成产品验证与交付标准；PRD、治理和融资材料随真实证据迭代。"),
]
for entry in whitepaper:
    add_slide_entry(doc, *entry)

add_h2(doc, "4.2 家庭教育大模型平台科技公司项目合作方案（10页）")
cooperation = [
    (1, "项目合作封面", "把讨论聚焦战略定位、落地路径、组织机制和合作模式。", "该材料是合作邀约而非完整商业计划。", "缺少具体合作对象、决策请求和下一步会议产出。", "封面后应增加合作目标、双方资源、决策事项与试点周期。"),
    (2, "背景与愿景", "提出平台赛道、数据驱动和跨赛道复用，并追求全链路自动化。", "愿景强，能激发合作方对平台潜力的理解。", "“替代传统教培”“百亿估值”“黑灯工厂”均过度确定。", "对外改为待验证假设；自动化目标改为人机协同和可审计运营。"),
    (3, "赛道价值判断", "从平台定位、估值逻辑和跨赛道复用论证价值。", "三个维度覆盖市场、资本与扩张。", "没有市场证据、可比公司、合规条件和数据权属基础。", "建立证据清单；未完成验证前不对外使用百亿或千亿估值表述。"),
    (4, "商业闭环目标", "用自动获客、成交、服务和风险控制定义“黑灯工厂”。", "追求低边际成本和流程自动化有经营价值。", "家庭教育不能无人化；“不会出现合规爆雷”是不应作出的保证。", "定义自动化率与人工升级率，并把合规表达改为持续治理目标。"),
    (5, "三大落地板块", "复用教培SaaS，把80%资源放在知识库、Agent和用户裂变。", "避免从零研发是正确的资源策略。", "裂变体系优先级可能高于结果验证；技术与知识库被合并也不够清晰。", "资源先投产品与交付闭环，再投裂变；技术、知识、数据治理分别设负责人。"),
    (6, "一年节奏", "用20%底座、80%用户端和一年闭环约束执行。", "强制聚焦有利于避免长期基础建设。", "百分比并非资源测算，缺少季度闸门和失败退出条件。", "按30/90/180/365天设置里程碑，资源比例随证据动态调整。"),
    (7, "组织与资本", "建议独立公司、资本人才同步和提前储备，并强调上市潜力。", "意识到组织载体与资源同步的重要性。", "资本关注和上市潜力未举证；新公司资产边界未定义。", "先完成资产、IP、数据、人员和客户迁移方案，再启动融资叙事。"),
    (8, "核心团队", "要求懂产品、用户、数据、心智的复合人才和All In。", "强调创业心态与决策效率，适合早期核心队伍。", "“All In”不能替代岗位能力、劳动安排与健康边界。", "用角色责任、决策权、90天产出和冲突机制定义核心团队。"),
    (9, "合作模式", "提出快速适配、角色分工、资源储备和业务扩张路径。", "路径顺序合理，适合合作启动。", "缺少费用、IP、数据控制、验收、终止和争议机制。", "试点协议至少明确范围、RACI、里程碑、预算、数据、IP和退出条款。"),
    (10, "合作结束页", "以共建新一代家庭教育入口收束。", "情绪完整，但行动请求不够具体。", "合作方无法从本页判断应立即做什么。", "改为明确30天试点决策、负责人、资源投入和下次评审日期。"),
]
for entry in cooperation:
    add_slide_entry(doc, *entry)

add_h2(doc, "4.3 榜样教育新商业模式对外宣发PPT（20页）")
promo = [
    (1, "新商业模式封面", "用孩子问题、父母成长和家庭改变定义新模式。", "一句话同时覆盖入口、价值和结果，品牌逻辑清晰。", "标题叫对外宣发，但内容实际更像产品与经营蓝图。", "对外版应减少内部机制和类比，保留用户价值与可信证据。"),
    (2, "从交易到关系", "比较经营对象、增长、交付、资产和定位的旧新逻辑。", "这是五份材料中最完整的商业模式转型表。", "“中国父母成长生态平台”仍应视为远景定位。", "把本页作为纲领总表，并在每一行配置年度指标。"),
    (3, "商业模式内核图", "组合拼多多增长、字节分发、海底捞服务和教育长期陪伴。", "类比帮助理解增长、算法、体验和交付四种能力。", "品牌借用过多可能掩盖自己的独特因果链，也可能引发误解。", "内部保留类比；对外改成自有“家庭成长操作系统”及证据链。"),
    (4, "四种成熟能力", "将外部类比转化为裂变、画像、服务与陪伴的对应设计。", "能把抽象借鉴落到业务模块。", "“越使用越懂家庭”必须受数据边界和模型评估约束。", "每种能力设置用户价值、风险边界、责任人和验证指标。"),
    (5, "小程序UI总览", "展示首页、测评、AI诊断、成长方案、陪跑和会员中心。", "完整呈现从入口到留存的产品愿景。", "首版功能过多，首页同时承担内容、任务、服务和商城，MVP风险高。", "首版只保留测评、任务、反馈、报告和人工服务入口。"),
    (6, "增长优化UI", "增加体检、每日任务、孩子端、排行榜、报告海报和AI管家。", "业务闭环清楚，游戏化可提升执行。", "儿童排行榜和公开成长结果可能制造压力、比较和隐私风险。", "默认关闭公开排名；奖励个人进步而非家庭间比较，分享需单独授权。"),
    (7, "用户成长旅程", "从触发、觉醒、行动、改变、长期到传播设计路径。", "把家长心理、平台动作、产品和指标对齐，适合产品蓝图。", "“温和呈现家庭互动影响”需要避免把问题归咎于父母。", "用非评判式语言呈现影响因素，并允许用户不同意或重新评估。"),
    (8, "分享裂变商城UI", "用邀请、拼团、积分商城和会员转介绍形成商业闭环。", "弱现金激励、强成长权益的方向优于传统分销。", "拼团与邀请奖励仍可能产生社交压力和分销观感。", "坚持单层奖励、清晰价格、无强制分享，并建立反作弊与投诉机制。"),
    (9, "裂变原则", "从现金佣金、卖课话术和代理身份转向成长权益与价值传播。", "风险意识较强，是裂变机制的治理页。", "需要再增加隐私、未成年人、广告真实性和奖励上限。", "把本页升级为增长合规准则并纳入产品验收。"),
    (10, "名师咨询与线下沙龙UI", "展示专家发现、预约、咨询、活动报名和复购路径。", "线上线下结合能提升信任和高客单转化。", "专家资质、儿童接触、录音录像、现场安全和投诉流程未体现。", "上线前建立专家准入、背景核验、服务边界、记录规则和应急预案。"),
    (11, "线上线下服务闭环", "用咨询、沙龙、城市活动、社群和档案形成长期服务。", "平台价值与收入方式对齐，适合服务产品规划。", "不同服务的责任、质量标准和毛利差异较大。", "先验证一种高信任服务，再复制；每类服务独立核算质量与经济性。"),
    (12, "用户社区UI", "展示打卡、成果、求助、评论、同城圈和身份等级。", "社区可以沉淀行动与口碑，降低育儿孤独感。", "家庭求助和儿童照片属于高敏感内容，存在暴露、误导和伤害风险。", "建立匿名选项、审核、举报、危机升级、未成年人影像保护和内容保留规则。"),
    (13, "社区经营价值", "将社区定义为成长结果传播场，而非吐槽广场。", "目标清楚，连接留存、信任、履约和传播。", "过度强调正面成果可能压制真实困难和负面反馈。", "允许求助与异议，区分支持社区、案例展示和营销内容。"),
    (14, "客户后台UI", "把会员、服务、订单、家庭档案和客服集中到家庭账户。", "这是把一次购买变为长期关系的关键界面。", "账户聚合增加权限错误和敏感信息泄露的影响面。", "家庭成员分权限、敏感字段遮蔽、操作留痕和导出删除必须进入设计。"),
    (15, "家庭账户经营", "解释会员身份、服务进度、订单、档案、活动和客服的经营价值。", "完整回答为什么“我的”不是普通订单中心。", "“每次服务不从零开始”必须与用户控制历史数据的权利平衡。", "用户可查看、纠正、下载和删除数据；新用途需重新说明与授权。"),
    (16, "第一个100万家庭", "提出内容/IP、老用户裂变、城市社群、机构合作和少量投放的五级增长结构。", "降低投流依赖的方向与白皮书一致。", "占比30/30/20/15/5没有实证，100万会显著放大交付和治理压力。", "先以1000、1万家庭阶段验证渠道质量，再决定规模和占比。"),
    (17, "产品与收入层级", "从免费测评到挑战、90天计划、会员、服务和生态形成收入阶梯。", "产品、用户价值、收入与战略作用对应完整。", "免费测评可能被滥用于制造焦虑，生态层不应早于供给治理。", "测评提供可操作价值且不夸大；会员成立后再开放生态。"),
    (18, "三年平台路线", "先用户资产化，再服务平台化，最后生态化。", "阶段顺序合理，并明确每年暂不做什么。", "“用户资产化”措辞可能造成数据所有权误解。", "改为“用户关系与授权数据可信化”，并以退出权和安全为前提。"),
    (19, "经营仪表盘", "把线索、成交、销售和复购指标升级为用户资产指标。", "覆盖获客、激活、活跃、结果、推荐和LTV，经营逻辑完整。", "缺少安全、投诉、人工效率、毛利和授权质量指标。", "形成平衡计分卡：增长、结果、经济、信任安全四类指标同时过关。"),
    (20, "对外收束", "把成熟课程、案例和服务升级为连接百万家庭的平台。", "愿景高度与前文一致，适合品牌收尾。", "“百万家庭”缺少当前基线和时间表。", "对外用“逐步连接更多家庭”，对内用阶段闸门管理规模目标。"),
]
for entry in promo:
    add_slide_entry(doc, *entry)

add_h2(doc, "4.4 榜样科技创业合伙人股权架构设计（14页）")
equity = [
    (1, "股权架构封面", "把股权定义为筛选、绑定和激励长期价值创造者的制度。", "目的清楚，避免把股权当普遍福利。", "仍应说明这是计划稿，不是授予承诺。", "所有沟通标注“方案讨论稿，以正式协议与审批为准”。"),
    (2, "创业合伙人制度", "分少数关键合伙人、多数核心骨干和4年成熟三个层次。", "分层有利于控制治理复杂度。", "“价值观绑定”若无明确标准，容易变成主观裁量。", "成熟条件以可观察行为和岗位结果为主，价值观采用有证据的行为标准。"),
    (3, "四大原则", "控制权、风险收益、贡献增量和分层激励构成设计原则。", "框架完整，可作为后续协议评审清单。", "只强调创始人控制，尚未体现激励对象的基本保护。", "增加透明、公平、可理解和可执行四项参与者保护原则。"),
    (4, "60/20/20架构", "创始团队60%、人才20%、资产与融资20%。", "为控制、激励和融资预留了结构空间。", "资产注入10%的估值、未来融资方式和稀释场景未说明。", "先做资产估值和三轮融资情景表，再决定比例；比例不是先验答案。"),
    (5, "人才双轨", "对比100%薪酬职业经理人与降薪换权益创业合伙人。", "用风险偏好区分合作方式，避免一刀切。", "本页叫双轨但下一页实际有B方案，叙事不一致；降薪需符合劳动安排。", "统一为A/B/C三轨，明确自愿选择、书面说明和转换窗口。"),
    (6, "三方案与创业系数", "用放弃现金×系数说明A、B、C风险收益差异。", "公式便于沟通风险与潜在回报。", "系数1/3/5缺少估值、概率、税务和稀释依据，不能直接定价。", "把公式改成场景测算工具，最终授予用估值和协议确定。"),
    (7, "六项机制", "覆盖换算、选择、成熟、持股平台、控制权和回购。", "已从比例讨论进入制度设计，是正确方向。", "每项仍是原则级，缺少审批主体、通知、争议和信息披露。", "形成权益政策、授予通知书、平台协议和退出流程四层文件体系。"),
    (8, "四年成熟", "设置12、24、36、48个月25%递进并叠加业绩条件。", "能降低人走股留和短期套利。", "“时间到了不一定成熟”若完全由董事会裁量，会削弱激励可信度。", "事先写明客观业绩条件、评估周期、补救期和争议处理。"),
    (9, "持股平台", "创始团队通过GP集中表决权，合伙人作为LP，大多数人才使用期权或虚拟股。", "治理结构简洁，便于融资和工商管理。", "具体载体、税负、登记、穿透和受益权需按当地规则审查。", "不得把“员工原则上不直接工商登记”当作通用法律结论，先做专项设计。"),
    (10, "三权边界", "区分经济收益、经营决策和董事权。", "有助于避免授予对象误解权利内容。", "单方面强调控制可能损害长期伙伴信任。", "协议同步明确知情权、收益权、稀释、关联交易回避和争议救济。"),
    (11, "融资稀释", "解释持股比例下降但权益价值可能增长。", "提前进行稀释教育是必要的。", "价值增长不是必然；期权池扩充、优先权和低估值融资会产生不同影响。", "提供至少三种融资情景和完全摊薄口径，禁止只讲乐观结果。"),
    (12, "退出与回购", "区分Good Leaver与Bad Leaver，并设计不同回购价格。", "退出机制是激励方案可执行的核心。", "“低价回购”必须有明确触发、证据、程序和适用性审查。", "定义事件、通知、申辩、估值日、支付方式和争议解决，避免模糊惩罚。"),
    (13, "实施路线", "两周定架构、四周定文件、2-3月试运行、4-12月扩展至50人。", "有阶段、有样本、有扩展节奏。", "50人、3000万和5-8名合伙人是未经经营计划验证的输入。", "先完成资本与人才预算，再决定池子规模；法律文件进度以关键决策齐备为前提。"),
    (14, "四机制收束", "用资格、成熟、回购和控制权四机制总结制度。", "总结准确，适合成为授予沟通首页。", "右上角存在文字重叠；且缺少税务、稀释和信息披露三项。", "修复版式，并增加税务提示、完全摊薄示例和年度权益报告。"),
]
for entry in equity:
    add_slide_entry(doc, *entry)

add_major_heading(doc, "结论与管理层决策建议")
add_callout(doc, "最终判断", "五份材料已经形成一致的战略方向：从卖课转向经营家庭成长，从一次交易转向长期关系，从名师依赖转向标准化交付，从数据记录转向可信数据闭环，从AI概念转向场景协同，从个人激励转向制度化合伙。但项目仍处于“方向基本一致、证据和治理尚未完备”的阶段。", PALE_GOLD, GOLD)

decisions = [
    ("决策一", "批准“家庭成长数据科技与服务平台建设者”作为内部战略定位；对外暂不宣称已经成为平台。"),
    ("决策二", "批准单一核心产品MVP：一个高频问题+家长训练营+AI陪练+成长报告，冻结非核心功能。"),
    ("决策三", "批准人机协同原则，否决无人化家庭教育服务和绝对合规承诺。"),
    ("决策四", "批准数据最小化、未成年人保护、人工升级和可撤回授权作为产品上线前置条件。"),
    ("决策五", "批准30/90/180/365天阶段闸门；100万家庭、跨赛道复制和全国扩张均须后置。"),
    ("决策六", "原则同意分层人才与4年成熟机制，但所有比例、系数、持股平台和回购条款须专项审查后实施。"),
]
for title, text in decisions:
    add_h2(doc, title)
    add_para(doc, text)

add_para(doc, "本文件建议作为下一阶段的母纲：产品PRD、交付SOP、数据治理制度、AI安全规范、合作协议、融资BP和股权激励文件都应与本纲领保持一致；出现冲突时，以“家庭真实改变、用户信任安全、可验证经营证据”三项原则为最高裁决标准。", bold=True, color=NAVY, before=8, after=0)

# Core document properties.
props = doc.core_properties
props.title = "榜样教育项目纲领与五份PPT逐页解读"
props.subject = "榜样教育战略、产品、数据AI、组织治理与股权机制汇总"
props.author = "Codex"
props.keywords = "榜样教育, 家庭教育, AI成长平台, 项目纲领, 逐页解读"
props.comments = "基于五份内部PPT整理；内部讨论稿。"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(str(OUT))
